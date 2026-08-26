"""Word and PDF export for generated ShadowIntel reports."""
from __future__ import annotations
import io


def report_to_docx(content: dict) -> io.BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading(content.get('title', 'Investigation Report'), level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Report ID: {content.get('id', '—')}    Generated: {content.get('generated', '—')}").italic = True

    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(content.get('executive_summary', ''))

    findings = content.get('system_findings', {})

    doc.add_heading('Key Entities', level=1)
    entities = findings.get('key_entities', [])
    if entities:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'ID', 'Label', 'Type', 'Priority'
        for entity in entities:
            row = table.add_row().cells
            row[0].text = str(entity.get('id', ''))
            row[1].text = str(entity.get('label', ''))
            row[2].text = str(entity.get('type', ''))
            row[3].text = str(entity.get('priority', ''))
    else:
        doc.add_paragraph('No key entities recorded.')

    doc.add_heading('Important Relationships', level=1)
    relationships = findings.get('important_relationships', [])
    if relationships:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, label in enumerate(['ID', 'Type', 'Source', 'Target', 'Evidence']):
            hdr[i].text = label
        for rel in relationships:
            row = table.add_row().cells
            row[0].text = str(rel.get('id', ''))
            row[1].text = str(rel.get('type', ''))
            row[2].text = str(rel.get('source', ''))
            row[3].text = str(rel.get('target', ''))
            row[4].text = str(rel.get('evidence_id', '') or 'Not recorded')
    else:
        doc.add_paragraph('No relationships recorded.')

    doc.add_heading('Cross-Source Correlations', level=1)
    correlations = findings.get('cross_source_correlations', [])
    if correlations:
        for corr in correlations:
            doc.add_paragraph(corr.get('explanation', ''), style='List Bullet')
    else:
        doc.add_paragraph('No observed cross-source correlations.')

    doc.add_heading('Risk / Anomaly Indicators', level=1)
    anomalies = findings.get('risk_anomalies', [])
    if anomalies:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Type', 'Risk / Severity', 'Explanation'
        for alert in anomalies:
            row = table.add_row().cells
            row[0].text = str(alert.get('type', ''))
            row[1].text = str(alert.get('risk_level', alert.get('severity', '')))
            row[2].text = str(alert.get('explanation', ''))
    else:
        doc.add_paragraph('No anomaly indicators recorded.')

    doc.add_heading('Timeline Highlights', level=1)
    timeline = findings.get('timeline_highlights', [])
    if timeline:
        for event in timeline:
            doc.add_paragraph(f"{event.get('timestamp', '')} — {event.get('title', '')} ({event.get('type', '')})", style='List Bullet')
    else:
        doc.add_paragraph('No timeline highlights recorded.')

    doc.add_heading('Source Evidence & Integrity', level=1)
    source_evidence = content.get('source_evidence', {})
    doc.add_paragraph(f"Verified: {source_evidence.get('verified_count', 0)}    Issues: {source_evidence.get('issue_count', 0)}")
    integrity = source_evidence.get('integrity', [])
    if integrity:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Evidence ID', 'Source', 'Status'
        for item in integrity[:30]:
            row = table.add_row().cells
            row[0].text = str(item.get('id', ''))
            row[1].text = str(item.get('source', ''))
            row[2].text = str(item.get('status', ''))

    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(content.get('methodology', ''))

    doc.add_heading('Responsible-AI Notice', level=1)
    disclaimer_paragraph = doc.add_paragraph(content.get('disclaimer', ''))
    if disclaimer_paragraph.runs:
        disclaimer_paragraph.runs[0].italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def report_to_pdf(content: dict) -> io.BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm, leftMargin=2 * cm, rightMargin=2 * cm)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(content.get('title', 'Investigation Report'), styles['Title']))
    story.append(Paragraph(f"Report ID: {content.get('id', '—')} &nbsp;&nbsp; Generated: {content.get('generated', '—')}", styles['Normal']))
    story.append(Spacer(1, 12))

    story.append(Paragraph('Executive Summary', styles['Heading1']))
    story.append(Paragraph(content.get('executive_summary', ''), styles['Normal']))
    story.append(Spacer(1, 10))

    findings = content.get('system_findings', {})

    def table_block(heading, rows, header):
        story.append(Paragraph(heading, styles['Heading1']))
        if rows:
            data = [header] + rows
            table = Table(data, hAlign='LEFT')
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#b3372c')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
        else:
            story.append(Paragraph('No data recorded.', styles['Normal']))
        story.append(Spacer(1, 10))

    entities = findings.get('key_entities', [])
    table_block('Key Entities', [[e.get('id', ''), e.get('label', ''), e.get('type', ''), str(e.get('priority', ''))] for e in entities], ['ID', 'Label', 'Type', 'Priority'])

    relationships = findings.get('important_relationships', [])
    table_block('Important Relationships', [[r.get('id', ''), r.get('type', ''), r.get('source', ''), r.get('target', ''), r.get('evidence_id', '') or 'Not recorded'] for r in relationships], ['ID', 'Type', 'Source', 'Target', 'Evidence'])

    correlations = findings.get('cross_source_correlations', [])
    story.append(Paragraph('Cross-Source Correlations', styles['Heading1']))
    if correlations:
        items = [ListItem(Paragraph(c.get('explanation', ''), styles['Normal'])) for c in correlations]
        story.append(ListFlowable(items, bulletType='bullet'))
    else:
        story.append(Paragraph('No observed cross-source correlations.', styles['Normal']))
    story.append(Spacer(1, 10))

    anomalies = findings.get('risk_anomalies', [])
    table_block('Risk / Anomaly Indicators', [[a.get('type', ''), a.get('risk_level', a.get('severity', '')), a.get('explanation', '')] for a in anomalies], ['Type', 'Risk', 'Explanation'])

    timeline = findings.get('timeline_highlights', [])
    story.append(Paragraph('Timeline Highlights', styles['Heading1']))
    if timeline:
        items = [ListItem(Paragraph(f"{e.get('timestamp', '')} — {e.get('title', '')} ({e.get('type', '')})", styles['Normal'])) for e in timeline]
        story.append(ListFlowable(items, bulletType='bullet'))
    else:
        story.append(Paragraph('No timeline highlights recorded.', styles['Normal']))
    story.append(Spacer(1, 10))

    source_evidence = content.get('source_evidence', {})
    story.append(Paragraph('Source Evidence & Integrity', styles['Heading1']))
    story.append(Paragraph(f"Verified: {source_evidence.get('verified_count', 0)} &nbsp;&nbsp; Issues: {source_evidence.get('issue_count', 0)}", styles['Normal']))
    integrity = source_evidence.get('integrity', [])[:30]
    if integrity:
        data = [['Evidence ID', 'Source', 'Status']] + [[i.get('id', ''), i.get('source', ''), i.get('status', '')] for i in integrity]
        table = Table(data, hAlign='LEFT')
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#b3372c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(table)
    story.append(Spacer(1, 10))

    story.append(Paragraph('Methodology', styles['Heading1']))
    story.append(Paragraph(content.get('methodology', ''), styles['Normal']))
    story.append(Spacer(1, 10))

    story.append(Paragraph('Responsible-AI Notice', styles['Heading1']))
    italic_style = ParagraphStyle('Disclaimer', parent=styles['Normal'], fontName='Helvetica-Oblique')
    story.append(Paragraph(content.get('disclaimer', ''), italic_style))

    doc.build(story)
    buffer.seek(0)
    return buffer